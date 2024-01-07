import os
import openai
import fitz  # PyMuPDF
from docx import Document
import pytesseract
from PIL import Image
from dotenv import load_dotenv
import time

# Change the current working directory to the script's directory
os.chdir(os.path.dirname(__file__))

# Load environment variables from .env file
load_dotenv()

# Initialize OpenAI client
client = openai.Client(api_key=os.getenv("OPENAI_API_KEY"))

# Set variables
logic_to_search = "Is this text about District 73 Toastmasters?"







# Set functions
def extract_text_from_pdf(file_path):
    try:
        text = ""
        with fitz.open(file_path) as doc:
            if not doc.page_count:  # Check if the document has any pages
                return None
            num_pages = min(3, doc.page_count)
            for page_num in range(num_pages):
                page = doc.load_page(page_num)
                page_text = page.get_text().strip()
                if not page_text:  # Use OCR if no text is found
                    pix = page.get_pixmap()
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    page_text = pytesseract.image_to_string(img)
                text += page_text + '\n'
        return text.strip() if text else None
    except Exception as e:
        return f"Error: {e}"

def extract_text_from_docx(file_path):
    try:
        text = ""
        with open(file_path, 'rb') as f:
            try:
                doc = Document(f)
            except ValueError:
                return None  # Return None if the document is corrupt or not a valid DOCX
            for para in doc.paragraphs:
                text += para.text + '\n'
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + ' \t'
                    text += '\n'
        return text.strip()
    except Exception as e:
        return f"Error: {e}"

def is_about_topic(text, file_path):
    if text:
        try:
            chat_completion = client.chat.completions.create(
                messages=[{"role": "user", "content": f"{logic_to_search} \n\n {text}"}],
                model="gpt-3.5-turbo"
            )
            response = chat_completion.choices[0].message.content

            # Write the response to a file
            response_file = file_path.replace(source_folder, output_folder) + "_api_response.txt"
            with open(response_file, 'w') as file:
                file.write(response)
                file.write("\n\n")
                file.write(text)

            return "yes" in response.lower()
        except Exception as e:
            print(f"Error in OpenAI API call for file {file_path}: {e}")
            return None
    else:
        return None

source_folder = 'input'
output_folder = 'output'
topic_folder = 'topic'
not_topic_folder = 'not-topic'

os.makedirs(source_folder, exist_ok=True)
os.makedirs(output_folder, exist_ok=True)
os.makedirs(topic_folder, exist_ok=True)
os.makedirs(not_topic_folder, exist_ok=True)

for root, dirs, files in os.walk(source_folder):
    for file in files:
        if file.endswith('.pdf') or file.endswith('.docx'):
            file_path = os.path.join(root, file)

            if file_path.endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif file_path.endswith('.docx'):
                text = extract_text_from_docx(file_path)
            else:
                text = None

            result = is_about_topic(text, file_path)
            if result is None:
                print(f"Skipping sorting of {file} due to API error.")
                continue
            
            # Add a delay before renaming
            time.sleep(1)

            if result:
                os.rename(file_path, os.path.join(topic_folder, file))
                print(f"Moved {file} to topic folder.")
            else:
                os.rename(file_path, os.path.join(not_topic_folder, file))
                print(f"Moved {file} to not-topic folder.")
